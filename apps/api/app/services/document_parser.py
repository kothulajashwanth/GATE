"""Document Parser Service for ExamShield AI.
Parses PDF, DOCX, TXT, and XLSX files to extract structured exam questions.
"""

import io
import re
from dataclasses import dataclass, field
from typing import Any

from openpyxl import load_workbook
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.errors import ValidationError
from app.db.models.question import Difficulty, Question, QuestionType


@dataclass
class ParsedQuestion:
    raw_number: str | None
    text: str
    question_type: QuestionType
    options: list[str]
    correct_answers: list[str]
    marks: int
    negative_marks: float
    difficulty: Difficulty
    explanation: str | None
    subject_code: str | None
    topic_name: str | None
    status: str  # "valid", "review_required", "duplicate"
    is_duplicate: bool = False
    duplicate_of_id: str | None = None
    failure_reason: str | None = None


@dataclass
class ExtractionResult:
    total: int = 0
    valid_count: int = 0
    review_count: int = 0
    failed_count: int = 0
    ocr_used: bool = False
    ocr_required: bool = False
    questions: list[dict[str, Any]] = field(default_factory=list)
    failed_questions: list[dict[str, Any]] = field(default_factory=list)


def normalize_text(text: str) -> str:
    """Normalize text for duplicate checking (lowercase, remove punctuation & space)."""
    return re.sub(r'[^a-z0-9]', '', text.lower())


def extract_pdf_text(file_bytes: bytes) -> tuple[str, bool]:
    """Extract text from PDF file. Returns (extracted_text, ocr_required)."""
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages_text.append(t)
        full_text = "\n".join(pages_text).strip()
        if len(full_text) < 20:
            return "", True
        return full_text, False
    except Exception:
        return "", True


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from DOCX document."""
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_txt = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_txt:
                    paragraphs.append(row_txt)
        return "\n".join(paragraphs)
    except Exception:
        # Fallback string decoder if docx module not available
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""


def extract_txt_text(file_bytes: bytes) -> str:
    """Extract text from TXT file trying UTF-8 and fallback encodings."""
    for enc in ["utf-8", "latin-1", "utf-16", "cp1252"]:
        try:
            return file_bytes.decode(enc).strip()
        except Exception:
            continue
    return file_bytes.decode("utf-8", errors="ignore").strip()


def parse_text_to_questions(text: str) -> list[dict[str, Any]]:
    """Parse raw text blocks into structured question dictionaries."""
    raw_blocks = re.split(r'\n(?=(?:Q\d+[\.:\)]|\d+[\.:\)]|Question\s+\d+[\.:\)]))', text.strip(), flags=re.IGNORECASE)

    extracted: list[dict[str, Any]] = []

    for idx, block in enumerate(raw_blocks):
        if not block.strip():
            continue

        lines = [l.strip() for l in block.split('\n') if l.strip()]
        if not lines:
            continue

        title_line = lines[0]
        # Match question header: Q1. / 1. / Question 1:
        num_match = re.match(r'^(?:Q\d+[\.:\)]|\d+[\.:\)]|Question\s+\d+[\.:\)])\s*(.*)', title_line, re.IGNORECASE)
        question_text = num_match.group(1).strip() if num_match else title_line

        if not question_text:
            extracted.append({
                "raw_data": block,
                "reason": "MISSING_QUESTION_TEXT",
                "is_failed": True,
            })
            continue

        options: list[str] = []
        answers: list[str] = []
        marks = 1
        negative_marks = 0.0
        difficulty = Difficulty.MEDIUM
        q_type = QuestionType.MCQ
        explanation = None
        subject_code = None
        topic_name = None

        for line in lines[1:]:
            # Match Options: A. Text, B) Text, a) Text, (i) Text
            opt_match = re.match(r'^(?:[A-D]\.|\([A-D]\)|[A-D]\)|[a-d]\.|\([a-d]\)|\([i|v|x]+\))\s*(.*)', line, re.IGNORECASE)
            if opt_match:
                opt_txt = opt_match.group(1).strip()
                if opt_txt:
                    options.append(opt_txt)
                continue

            # Match Correct Answer: Answer: B, Ans: B, Correct: B, Key: B
            ans_match = re.match(r'^(?:Answer|Correct|Key|Ans)[\s:]*([A-D]|True|False|[0-9a-zA-Z\s,]+)', line, re.IGNORECASE)
            if ans_match:
                ans_val = ans_match.group(1).strip().upper()
                if ans_val in ["A", "B", "C", "D"]:
                    answers.append(ans_val)
                elif ans_val in ["TRUE", "FALSE"]:
                    answers.append(ans_val)
                    q_type = QuestionType.TRUE_FALSE
                else:
                    answers.append(ans_val)
                continue

            # Match Marks: Marks: 2
            marks_match = re.match(r'^Marks[\s:]*(\d+)', line, re.IGNORECASE)
            if marks_match:
                try:
                    marks = int(marks_match.group(1))
                except ValueError:
                    pass
                continue

            # Match Difficulty: EASY/MEDIUM/HARD
            diff_match = re.match(r'^Difficulty[\s:]*(EASY|MEDIUM|HARD)', line, re.IGNORECASE)
            if diff_match:
                d_str = diff_match.group(1).upper()
                if d_str in Difficulty.__members__:
                    difficulty = Difficulty[d_str]
                continue

            # Match Explanation:
            exp_match = re.match(r'^(?:Explanation|Solution)[\s:]*(.*)', line, re.IGNORECASE)
            if exp_match:
                explanation = exp_match.group(1).strip()
                continue

            # Match Subject:
            sub_match = re.match(r'^Subject[\s:]*(.*)', line, re.IGNORECASE)
            if sub_match:
                subject_code = sub_match.group(1).strip()
                continue

            # Match Topic:
            top_match = re.match(r'^Topic[\s:]*(.*)', line, re.IGNORECASE)
            if top_match:
                topic_name = top_match.group(1).strip()
                continue

        # Determine Question Type and Validation Status
        status = "valid"
        reason = None

        if len(options) == 0:
            if answers and answers[0] in ["TRUE", "FALSE"]:
                options = ["True", "False"]
                q_type = QuestionType.TRUE_FALSE
            else:
                q_type = QuestionType.SHORT_ANSWER
        elif len(options) < 2:
            status = "failed"
            reason = "INVALID_OPTION_FORMAT"
            extracted.append({
                "raw_data": block,
                "reason": reason,
                "is_failed": True,
            })
            continue

        if not answers:
            status = "review_required"
            # Default to first option for preview editing convenience
            answers = ["A"] if options else [""]

        extracted.append({
            "number": str(idx + 1),
            "text": question_text,
            "question_type": q_type.value,
            "options": options,
            "correct_answers": answers,
            "marks": marks,
            "negative_marks": negative_marks,
            "difficulty": difficulty.value,
            "explanation": explanation,
            "subject_code": subject_code,
            "topic_name": topic_name,
            "status": status,
            "is_failed": False,
        })

    return extracted


def parse_excel_questions(file_bytes: bytes) -> list[dict[str, Any]]:
    """Parse XLSX files for questions."""
    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))

    if not rows or len(rows) < 2:
        return []

    questions = []
    for idx, row in enumerate(rows[1:]):
        if not row or not any(row):
            continue

        q_text = str(row[0]).strip() if len(row) > 0 and row[0] else ""
        if not q_text:
            questions.append({
                "raw_data": str(row),
                "reason": "MISSING_QUESTION_TEXT",
                "is_failed": True,
            })
            continue

        opt_a = str(row[1]).strip() if len(row) > 1 and row[1] else "Option A"
        opt_b = str(row[2]).strip() if len(row) > 2 and row[2] else "Option B"
        opt_c = str(row[3]).strip() if len(row) > 3 and row[3] else "Option C"
        opt_d = str(row[4]).strip() if len(row) > 4 and row[4] else "Option D"
        ans = str(row[5]).strip().upper() if len(row) > 5 and row[5] else "A"
        diff_str = str(row[6]).strip().upper() if len(row) > 6 and row[6] else "MEDIUM"
        marks_str = str(row[7]).strip() if len(row) > 7 and row[7] else "1"

        diff = Difficulty.MEDIUM
        if diff_str in Difficulty.__members__:
            diff = Difficulty[diff_str]

        try:
            marks = int(marks_str)
        except ValueError:
            marks = 1

        questions.append({
            "number": str(idx + 1),
            "text": q_text,
            "question_type": QuestionType.MCQ.value,
            "options": [opt_a, opt_b, opt_c, opt_d],
            "correct_answers": [ans if ans in ["A", "B", "C", "D"] else "A"],
            "marks": marks,
            "negative_marks": 0.0,
            "difficulty": diff.value,
            "explanation": "Imported from Excel file",
            "subject_code": None,
            "topic_name": None,
            "status": "valid",
            "is_failed": False,
        })

    return questions


async def process_uploaded_document(
    db: AsyncSession,
    file_bytes: bytes,
    filename: str,
) -> ExtractionResult:
    """Master pipeline to extract text, parse questions, validate, and flag duplicates."""
    ext = filename.lower().split('.')[-1]
    result = ExtractionResult()

    if ext in ['xlsx', 'xls']:
        raw_items = parse_excel_questions(file_bytes)
    elif ext == 'pdf':
        text, ocr_req = extract_pdf_text(file_bytes)
        result.ocr_required = ocr_req
        if ocr_req:
            return result
        raw_items = parse_text_to_questions(text)
    elif ext in ['docx', 'doc']:
        text = extract_docx_text(file_bytes)
        raw_items = parse_text_to_questions(text)
    else:  # txt or default
        text = extract_txt_text(file_bytes)
        raw_items = parse_text_to_questions(text)

    # Fetch existing questions from DB to check normalized duplicates
    existing_db_qs = (await db.execute(select(Question.id, Question.text).where(Question.deleted_at.is_(None)))).all()
    existing_norm_map = {normalize_text(q.text): str(q.id) for q in existing_db_qs if q.text}

    seen_in_batch: set[str] = set()

    for item in raw_items:
        result.total += 1
        if item.get("is_failed"):
            result.failed_count += 1
            result.failed_questions.append({
                "row": result.total,
                "raw_data": item.get("raw_data", ""),
                "reason": item.get("reason", "PARSER_ERROR"),
            })
            continue

        q_text = item["text"]
        norm = normalize_text(q_text)

        is_duplicate = False
        duplicate_of_id = None

        if norm in existing_norm_map:
            is_duplicate = True
            duplicate_of_id = existing_norm_map[norm]
        elif norm in seen_in_batch:
            is_duplicate = True

        seen_in_batch.add(norm)

        item["is_duplicate"] = is_duplicate
        item["duplicate_of_id"] = duplicate_of_id
        if is_duplicate:
            item["status"] = "review_required"

        if item["status"] == "valid":
            result.valid_count += 1
        else:
            result.review_count += 1

        result.questions.append(item)

    return result
